from docx import Document
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName,
    FileType, Disposition, ContentId)
import os
import datetime as date
import base64
import io
from audiorecorder import audiorecorder
from openai import OpenAI
import tempfile
from webapp.constants import *
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory

def create_interview_file(username: str, patient: str, messages: list[dict[str, str]]) -> Document:
    interview = Document()
    heading = interview.add_paragraph("User: " + username + ", ")
    currentDateAndTime = date.datetime.now()
    date_time = currentDateAndTime.strftime("%d-%m-%y__%H-%M")
    heading.add_run("Date: " + date_time + ", ")
    heading.add_run("Patient: " + patient)
    for message in messages:
        interview.add_paragraph(message["role"] + ": " + message["content"])
    
    return interview


def send_email(bio, EMAIL_TO_SEND, username, date_time, feedback):
    message = Mail(
        from_email = 'rutgers.aime@gmail.com',
        to_emails = EMAIL_TO_SEND,
        subject = "Conversation from " + username + " at time " + date_time,
        html_content = feedback)
    attachment = Attachment()
    encoded = base64.b64encode(bio.getvalue()).decode()
    attachment.file_content=FileContent(encoded)
    attachment.file_type = FileType('docx')
    attachment.file_name = FileName(username + "_" + date_time + ".docx")
    attachment.disposition = Disposition('attachment')
    attachment.content_id = ContentId('docx')
    message.attachment = attachment
    try:
        sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
        response = sg.send(message)
        print(response.status_code)
        print(response.body)
        print(response.headers)
    except: 
        print("ERROR ENCOUNTERED SENDING MESSAGE\n")

def transcribe_voice(voice_input, OPENAI_API_KEY):
    client = OpenAI()
    temp_audio_file = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
    voice_input.export(temp_audio_file.name, format="wav")
    with open(temp_audio_file.name, "rb") as file:
        transcription = client.audio.transcriptions.create(model="whisper-1", 
                                                    file=file, 
                                                    response_format="text")
    
    return transcription

def classify_input(messages: list[str], OPENAI_API_KEY: str) -> dict[str, bool]:
    llm = ChatOpenAI(api_key=OPENAI_API_KEY, model=MODEL, temperature=0.0)
    conversation = ConversationChain(llm=llm, memory=ConversationBufferMemory())
    with open(CLASSIFY_INPUT_PROMPT, "r", encoding="utf8") as classify:
        prompt_input = classify.read()
    for message in messages:
        message = message.rstrip() + " "
        prompt_input += message
    raw_classification = conversation.predict(input=prompt_input)
    classification = raw_classification.split()
    output = {}
    for label in CLASSIFY_INPUT_LABELS:
        if label in classification:
            output[label] = True
        else:
            output[label] = False
    return output